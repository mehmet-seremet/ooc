"""
Markdown to Word Converter for Genius OOC Documentation
This script converts the Markdown documentation to a formatted Word document
"""

import pandas as pd
from docx import Document
from docx.shared import Inches
import re
import os

def read_markdown_file(file_path):
    """Read markdown file and return content"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except FileNotFoundError:
        print(f"Error: File {file_path} not found")
        return None
    except Exception as e:
        print(f"Error reading file: {e}")
        return None

def parse_markdown_content(content):
    """Parse markdown content and extract structured data"""
    lines = content.split('\n')
    
    # Data structure to hold parsed content
    parsed_data = {
        'title': '',
        'sections': [],
        'current_section': None,
        'current_subsection': None,
        'current_content': []
    }
    
    for line in lines:
        line = line.strip()
        
        # Main title (# Title)
        if line.startswith('# ') and not parsed_data['title']:
            parsed_data['title'] = line[2:].strip()
        
        # Section headers (## Section)
        elif line.startswith('## '):
            # Save previous section if exists
            if parsed_data['current_section']:
                parsed_data['sections'].append({
                    'title': parsed_data['current_section'],
                    'subsections': parsed_data.get('subsections', []),
                    'content': '\n'.join(parsed_data['current_content'])
                })
            
            parsed_data['current_section'] = line[3:].strip()
            parsed_data['subsections'] = []
            parsed_data['current_content'] = []
        
        # Subsection headers (### Subsection)
        elif line.startswith('### '):
            if parsed_data['current_subsection']:
                parsed_data['subsections'].append({
                    'title': parsed_data['current_subsection'],
                    'content': '\n'.join(parsed_data['current_content'])
                })
                parsed_data['current_content'] = []
            
            parsed_data['current_subsection'] = line[4:].strip()
        
        # Regular content
        elif line:
            parsed_data['current_content'].append(line)
    
    # Add the last section
    if parsed_data['current_section']:
        if parsed_data['current_subsection']:
            parsed_data['subsections'].append({
                'title': parsed_data['current_subsection'],
                'content': '\n'.join(parsed_data['current_content'])
            })
        
        parsed_data['sections'].append({
            'title': parsed_data['current_section'],
            'subsections': parsed_data.get('subsections', []),
            'content': '\n'.join(parsed_data['current_content'])
        })
    
    return parsed_data

def create_dataframe_from_parsed_data(parsed_data):
    """Create a pandas DataFrame from parsed markdown data"""
    rows = []
    
    # Add title row
    rows.append({
        'Type': 'Title',
        'Level': 1,
        'Title': parsed_data['title'],
        'Content': '',
        'Section': '',
        'Subsection': ''
    })
    
    # Add sections and subsections
    for section in parsed_data['sections']:
        # Add section header
        rows.append({
            'Type': 'Section Header',
            'Level': 2,
            'Title': section['title'],
            'Content': '',
            'Section': section['title'],
            'Subsection': ''
        })
        
        # Add section content if exists
        if section['content'].strip():
            rows.append({
                'Type': 'Content',
                'Level': 2,
                'Title': '',
                'Content': section['content'],
                'Section': section['title'],
                'Subsection': ''
            })
        
        # Add subsections
        for subsection in section.get('subsections', []):
            rows.append({
                'Type': 'Subsection Header',
                'Level': 3,
                'Title': subsection['title'],
                'Content': '',
                'Section': section['title'],
                'Subsection': subsection['title']
            })
            
            if subsection['content'].strip():
                rows.append({
                    'Type': 'Content',
                    'Level': 3,
                    'Title': '',
                    'Content': subsection['content'],
                    'Section': section['title'],
                    'Subsection': subsection['title']
                })
    
    return pd.DataFrame(rows)

def format_content_for_word(content):
    """Format markdown content for Word document"""
    # Remove markdown formatting
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
    content = re.sub(r'\*(.*?)\*', r'\1', content)      # Italic
    content = re.sub(r'`(.*?)`', r'\1', content)        # Code
    content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)  # Links
    
    return content

def create_word_document(df, output_path):
    """Create Word document from DataFrame"""
    # Create new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process DataFrame and add content to document
    for index, row in df.iterrows():
        if row['Type'] == 'Title':
            # Use built-in Title style
            p = doc.add_paragraph(row['Title'], style='Title')
            
        elif row['Type'] == 'Section Header':
            # Use built-in Heading 1 style
            p = doc.add_paragraph(row['Title'], style='Heading 1')
            
        elif row['Type'] == 'Subsection Header':
            # Use built-in Heading 2 style
            p = doc.add_paragraph(row['Title'], style='Heading 2')
            
        elif row['Type'] == 'Content' and row['Content'].strip():
            # Format content
            formatted_content = format_content_for_word(row['Content'])
            
            # Handle lists and special formatting
            lines = formatted_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Handle bullet points
                if line.startswith('- '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('* '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                # Handle numbered lists
                elif re.match(r'^\d+\.', line):
                    p = doc.add_paragraph(line[3:].strip(), style='List Number')
                else:
                    p = doc.add_paragraph(line, style='Normal')
    
    # Save document
    try:
        doc.save(output_path)
        print(f"Word document successfully created: {output_path}")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

def main():
    """Main function to convert markdown to Word"""
    # File paths
    markdown_file = "genius_ooc_documentation.md"
    word_file = "genius_ooc_documentation.docx"
    
    print("Starting Markdown to Word conversion...")
    print(f"Reading markdown file: {markdown_file}")
    
    # Read markdown content
    content = read_markdown_file(markdown_file)
    if content is None:
        return
    
    print("Parsing markdown content...")
    
    # Parse markdown content
    parsed_data = parse_markdown_content(content)
    
    print(f"Found {len(parsed_data['sections'])} sections")
    
    # Create DataFrame
    df = create_dataframe_from_parsed_data(parsed_data)
    
    print(f"Created DataFrame with {len(df)} rows")
    print("\nDataFrame structure:")
    print(df[['Type', 'Level', 'Title', 'Section', 'Subsection']].head(10))
    
    # Save DataFrame to CSV for inspection
    csv_file = "documentation_structure.csv"
    df.to_csv(csv_file, index=False, encoding='utf-8')
    print(f"\nDataFrame saved to CSV: {csv_file}")
    
    print(f"\nCreating Word document: {word_file}")
    
    # Create Word document
    success = create_word_document(df, word_file)
    
    if success:
        print("\n✅ Conversion completed successfully!")
        print(f"📄 Markdown file: {markdown_file}")
        print(f"📊 CSV structure file: {csv_file}")
        print(f"📝 Word document: {word_file}")
    else:
        print("\n❌ Conversion failed!")

if __name__ == "__main__":
    main()